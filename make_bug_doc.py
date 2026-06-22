#!/usr/bin/env python3
"""Generate a formatted Word (.docx) version of the bug report."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

NAVY = RGBColor(0x1A, 0x1A, 0x2E)
BLUE = RGBColor(0x0F, 0x34, 0x60)
ACCENT = RGBColor(0x00, 0x7A, 0xCC)

doc = Document()

# Base style
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# Title
title = doc.add_heading('Battleship — Bugs Found & Fixes', level=0)
for run in title.runs:
    run.font.color.rgb = NAVY

sub = doc.add_paragraph('A record of the bugs uncovered during development and how each was resolved.')
sub.runs[0].italic = True
sub.runs[0].font.color.rgb = RGBColor(0x55, 0x5B, 0x6E)


def section(text):
    h = doc.add_heading(text, level=1)
    for r in h.runs:
        r.font.color.rgb = BLUE


def bug(num, name, items):
    """items: list of (label, text) tuples."""
    h = doc.add_heading(f'{num}. {name}', level=2)
    for r in h.runs:
        r.font.color.rgb = ACCENT
    for label, text in items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f'{label}: ')
        run.bold = True
        p.add_run(text)


section('Gameplay Bugs')

bug(1, "Couldn't attack the opponent", [
    ('Symptom', 'After Start Game, clicking the enemy board did nothing.'),
    ('Cause', 'The attack click listener was attached to your own board instead of the opponent\u2019s board.'),
    ('Fix', 'Attached the listener to the AI board and made it always-on, with turn/game-over checks handled inside the attack handler so the board can never become "dead."'),
])

bug(2, 'Board stopped responding after a miss', [
    ('Symptom', 'Attacks worked for a click or two, then stopped.'),
    ('Cause', 'After the AI\u2019s turn, only the player board was re-rendered, so the opponent board lost its click handlers.'),
    ('Fix', 'Re-render the opponent board when control returns to the player; the always-on listener makes this robust.'),
])

bug(3, 'Unfair turns', [
    ('Symptom', 'The player got an extra shot after a hit, but the AI never did.'),
    ('Cause', 'Inconsistent turn logic between player and AI.'),
    ('Fix', 'The AI now also earns another shot on a hit \u2014 rules are symmetric.'),
])

bug(4, 'AI wasted shots after sinking a ship', [
    ('Symptom', 'The AI kept firing around a ship it had already sunk.'),
    ('Cause', '"Target mode" wasn\u2019t cleared after a sink.'),
    ('Fix', 'The sink is now reported and the AI clears its target queue and resumes hunting.'),
])

section('Stability Bugs')

bug(5, 'AI turn crashed the game', [
    ('Cause', 'A variable was declared inside a loop but used outside it, throwing a ReferenceError.'),
    ('Fix', 'Declared the variable in the correct scope.'),
])

bug(6, 'AI could freeze', [
    ('Cause', 'The AI re-picked already-hit cells using a retry loop that could spin forever.'),
    ('Fix', 'Rewrote cell selection to build a list of un-attacked cells and pick from it \u2014 guaranteed to terminate.'),
])

bug(7, 'State leaked between games', [
    ('Cause', 'Ship data, game-over, turn state, and AI tracking weren\u2019t fully reset on a new game; ship-button listeners stacked up each replay.'),
    ('Fix', 'Setup now resets all state, and ship buttons use single-assignment handlers so listeners never accumulate.'),
])

bug(8, 'Hit counts overflowed', [
    ('Cause', 'A sunk ship could keep incrementing its hit total.'),
    ('Fix', 'Added a guard that stops counting once hits reach the ship length.'),
])

bug('8b', 'Stale AI timer caused a false "You Lost" after restart', [
    ('Symptom', 'Clicking "New Game" during the AI\u2019s 1-second delay could pop a "You Lost!" screen over the fresh setup a moment later.'),
    ('Cause', 'The pending AI-turn timer still fired after the board was reset, then attacked an empty board where the player had no ships \u2014 tripping the loss condition.'),
    ('Fix', 'The AI timer is now tracked and cancelled on reset, and the AI turn bails out early if no game is active. Verified with an automated test.'),
])

section('UI / UX Bugs')

bug(9, 'Page required scrolling', [
    ('Cause', 'Fixed sizing overflowed shorter screens.'),
    ('Fix', 'The board size is derived from the leftover viewport height, and the layout is locked to one screen.'),
])

bug(10, 'Orientation control', [
    ('Change', 'Replaced the rotate button with drag-to-place (sideways = horizontal, up/down = vertical), plus R key and right-click to flip; later removed the leftover button while keeping all of that functionality.'),
])

bug(11, "Edits didn't appear in the browser", [
    ('Cause', 'The dev server returned cached responses, serving stale files.'),
    ('Fix', 'Added a no-cache dev server and versioned asset URLs so the latest always loads.'),
])

section('How Fixes Were Verified')

p = doc.add_paragraph()
p.add_run('Logic tests (JavaScript engine, DOM-stubbed): ').bold = True
for line in [
    'Syntax/parse check passes',
    'Ship placement and validation (including off-board) are correct',
    'Player can win (all AI ships sunk, score 5/5)',
    'AI can win (repeated turns sink all player ships)',
    'Repeated attacks on the same cell don\u2019t double-count',
    'Hit counts never exceed ship length',
    'All element IDs referenced in JS exist in the HTML',
]:
    doc.add_paragraph(line, style='List Bullet')

p2 = doc.add_paragraph()
p2.add_run('Browser UI tests (Playwright, Chromium + Firefox): ').bold = True
for line in [
    'Places all ships through the UI and starts a game',
    'Player can win by clicking every AI ship cell',
    'Re-attacking the same cell is rejected',
    'R key rotates a ship to vertical placement',
    'Restarting during the AI delay does not cause a false loss',
]:
    doc.add_paragraph(line, style='List Bullet')

closing = doc.add_paragraph('All checks passed (10 browser tests across 2 engines, plus the logic suites).')
closing.runs[0].bold = True

out = 'Battleship_Bugs_Fixed.docx'
doc.save(out)
print('Saved', out)
